# streamlit_app.py
from datetime import date, datetime, timedelta

from docx import Document
from langchain_community.callbacks import get_openai_callback
from llm_utils.streamlit_common import hide_streamlit_branding
from llm_utils.text_format import convert_markdown_docx

import CV_Promoter_config.config as config
import CV_Promoter_config.instructions_config as instructions_config
import streamlit as st

# TODO fix streamlit implementation
from CV_Promoter.workflows import (
    AnnualReviewDrafter,
    NarrativePortfolioDrafter,
    RecommendationLetterDrafter,
)

this_year = datetime.now().year

# Define the available sections of interest
sections_of_interest = instructions_config.section_instructions.keys()
portfolio_sections = instructions_config.narrative_instructions.keys()


def show_CV_Promoter_page(template_location=config.TEMPLATE):
    """
    The `show_CV_Promoter_page` function in Python creates a Streamlit page for users to upload a
    Curriculum Vitae in Word format and generate drafts for Promotion Portfolio, Annual Review, and
    Recommendation Letter based on the uploaded CV.

    Args:
      template_location: The `template_location` parameter in the `show_CV_Promoter_page` function is
    used to specify the location of the template file that will be used for generating the output
    documents (e.g., draft narrative, recommendation letter). This parameter allows you to customize the
    template used for creating these documents
    """
    # page metadata
    st.set_page_config(page_title="CV Promoter", page_icon="🎓")
    # hide streamlit branding
    hide_streamlit_branding()

    # page content
    st.title("🎓 Promotion and Review 💼")
    st.markdown(
        """
    **Get help with promotion and tenure (P&T) or annual review materials.**

    Brought to you by the Anesthesiology Research, Informatics, and Perioperative Data Science teams in collaboration with Radiology Imaging Informatics, Clinicians, and Researchers.

    ---
    """
    )

    # Allow user to upload a Word cv
    uploaded_file = st.file_uploader("Choose a Curriculum Vitae in Microsoft Word format", type="docx")

    tab1, tab2, tab3 = st.tabs(["Promotion Portfolio", "Annual Review", "Recommendation Letter"])

    with tab2:
        st.write("Assistant for preparing Annual Review forms.")
        if uploaded_file is not None:
            cv_document = Document(uploaded_file)
            st.write("CV uploaded successfully!")
            # Create a dropdown for the user to select a section of interest
            section_of_interest = st.selectbox("Select a section of interest:", sections_of_interest)

            if st.button("Extract"):
                # submit prompt
                with st.spinner("Extracting. This may take a while..."):
                    with get_openai_callback() as response_meta:
                        submit_time = datetime.now()
                        workflow = AnnualReviewDrafter(section_of_interest, cv_document)
                        generated_response = workflow.process()
                        response_time = datetime.now()
                        form_docx_data = convert_markdown_docx(generated_response.content, template_location)

                if form_docx_data:
                    st.balloons()
                    st.write("Note that once you hit download, this form will reset.")

                    st.download_button(
                        label="Download form draft",
                        data=form_docx_data,
                        file_name="DRAFT_filled_form.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # correct MIME type for docx
                    )
        else:
            st.write("Please upload a Word Document (.docx) CV to proceed.")

    with tab1:
        st.write("Get a draft of a narrative section of your Promotion and Tenure portfolio.")
        if uploaded_file is not None:
            cv_document = Document(uploaded_file)
            st.write("CV uploaded successfully!")
            # Create a dropdown for the user to select a section of interest
            section_of_interest = st.selectbox("Select a section of interest:", portfolio_sections)

            selected_date = st.date_input(
                "What was your start date?",
                date.today() + timedelta(weeks=-260),
                min_value=datetime.now() + timedelta(weeks=-2600),
            )

            if selected_date:
                start_year = selected_date.year

            if st.button("Draft narrative"):
                # submit prompt
                with st.spinner("Drafting. This may take a while..."):
                    with get_openai_callback() as response_meta:
                        submit_time = datetime.now()
                        workflow = NarrativePortfolioDrafter(section_of_interest, selected_date, cv_document)
                        generated_response = workflow.process()
                        response_time = datetime.now()
                        portfolio_docx_data = convert_markdown_docx(
                            generated_response.content, template_location
                        )

                if portfolio_docx_data:
                    st.balloons()
                    st.write("Note that once you hit download, this form will reset.")

                    st.download_button(
                        label="Download narrative draft",
                        data=portfolio_docx_data,
                        file_name="DRAFT_portfolio_section.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # correct MIME type for docx
                    )

        else:
            st.write("Please upload a Word Document (.docx) CV to proceed.")

    with tab3:
        st.write("Get a draft recommendation letter for a colleague.")
        if uploaded_file is not None:
            cv_document = Document(uploaded_file)
            st.write("CV uploaded successfully!")
            # Create a dropdown for the user to select a section of interest
            areas_of_excellence = st.multiselect(
                "Select one or two areas of excellence. Primary focus will be on the first one selected.",
                portfolio_sections,
            )

            # Ensure that the user selects at most two sections
            if len(areas_of_excellence) > 2:
                st.warning("Please select at most two areas.")

            selected_date = st.date_input(
                "What was their start date?",
                date.today() + timedelta(weeks=-260),
                min_value=datetime.now() + timedelta(weeks=-2600),
            )

            st.write(
                "***Double check the date.*** It is most important that the year be correct. If using the calendar tool, the last date selected will be inserted into the blank."
            )

            if st.button("Draft letter"):
                # submit prompt
                with st.spinner("Drafting. This may take a while..."):
                    with get_openai_callback() as response_meta:
                        submit_time = datetime.now()
                        workflow = RecommendationLetterDrafter(
                            areas_of_excellence, selected_date, cv_document
                        )
                        generated_response = workflow.process()
                        response_time = datetime.now()
                        portfolio_docx_data = convert_markdown_docx(
                            generated_response.content, template_location
                        )

                if portfolio_docx_data:
                    st.balloons()
                    st.write("Note that once you hit download, this form will reset.")

                    st.download_button(
                        label="Download letter draft",
                        data=portfolio_docx_data,
                        file_name="DRAFT_recommendation_letter.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # correct MIME type for docx
                    )

        else:
            st.write("Please upload a Word Document (.docx) CV to proceed.")


if __name__ == "__main__":
    show_CV_Promoter_page()
